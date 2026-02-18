from io import BytesIO
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
import base64
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_horizontal_line(doc, length_in_inches, align):
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_before = 0
    p_format.space_after = 0

    # Si se define longitud, centramos la línea ajustando márgenes
    if length_in_inches is not None:
        total_width = 6.0  # ancho estimado del cuerpo de texto en pulgadas
        margin = max((total_width - length_in_inches) / 2, 0)
        p_format.left_indent = Inches(margin)
        p_format.right_indent = Inches(margin)
        if align == "center":
            margin = max((total_width - length_in_inches) / 2, 0)
            p_format.left_indent = Inches(margin)
            p_format.right_indent = Inches(margin)
        elif align == "right":
            p_format.left_indent = Inches(total_width - length_in_inches)
            p_format.right_indent = Inches(0)
        elif align == "left":
            p_format.left_indent = Inches(0)
            p_format.right_indent = Inches(total_width - length_in_inches)

    p_element = p._p
    pPr = p_element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # grosor
    bottom.set(qn('w:space'), '1')          # separación con texto
    bottom.set(qn('w:color'), '000000')        # color
    pBdr.append(bottom)
    pPr.append(pBdr)

def insertar_imagen(base64_img, ancho, doc):
        if base64_img.startswith('data:image'):
            base64_img = base64_img.split(',')[1]

        img_data = base64.b64decode(base64_img)
        stream = BytesIO(img_data)

        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(stream, width=Inches(ancho))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER