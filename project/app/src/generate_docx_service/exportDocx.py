import base64
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT ,WD_LINE_SPACING
from django.http import HttpResponse
from docx import Document
import datetime
from datetime import timezone
from app.src.generate_docx_service.parse_markdown_to_docx import markdown_to_docx
from app.src.generate_docx_service.utils import add_horizontal_line, insertar_imagen
from app.core.auth.firebase_config import db
from app.src.get_database_report_data import getDataDBWord


def ProcessDocx(request):
     task_id = request.session.get("task_id")

     AiMarkdown, graphic, calendarios, horas, now_str,  place_str, tabla_img = getDataDBWord(task_id)

     doc = Document()

     #Configuracion de estilos:
     style_h1 = doc.styles['Heading 1']
     style_h1.font.size = Pt(20)
     style_h1.font.bold = True
     style_h1.font.color.rgb = RGBColor(0, 0, 0)
     style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
     style_h1.paragraph_format.space_before = Pt(24)
     style_h1.paragraph_format.space_after = Pt(12)
     style_h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
     style_h1.paragraph_format.line_spacing = 1.2

     style_h2 = doc.styles['Heading 2']
     style_h2.font.name = 'Calibri'
     style_h2.font.bold = True
     style_h2.font.color.rgb = RGBColor(0, 0, 0)
     style_h2.font.size = Pt(16)
     style_h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
     style_h2.paragraph_format.space_before = Pt(24)
     style_h2.paragraph_format.space_after = Pt(12)
     style_h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
     style_h2.paragraph_format.line_spacing = 1.2

     style_h3 = doc.styles['Heading 3']
     style_h3.font.name = 'Calibri'
     style_h3.font.bold = True
     style_h3.font.color.rgb = RGBColor(0, 0, 0)
     style_h3.font.size = Pt(14)
     style_h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
     style_h3.paragraph_format.space_before = Pt(24)
     style_h3.paragraph_format.space_after = Pt(12)
     style_h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
     style_h3.paragraph_format.line_spacing = 1.2

     style_p = doc.styles['Normal']
     style_p.font.name = 'Calibri'
     style_p.font.size = Pt(11)
     style_p.font.color.rgb = RGBColor(0, 0, 0)
     style_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
     style_p.paragraph_format.line_spacing = 1.15

     # Portada
     doc.add_paragraph("\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n")  # espacio
     titulo = doc.add_heading("Análisis de Eventos", level=1)
     titulo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
     subtitulo = doc.add_paragraph(place_str)
     subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

     add_horizontal_line(doc, 2.5, 'right')

     fecha = doc.add_paragraph()
     fecha.add_run("Fecha: ").bold = True
     fecha.add_run(str(now_str or ""))
     fecha.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
     

     doc.add_page_break()
     if AiMarkdown:
        markdown_to_docx(AiMarkdown, doc)
     doc.add_page_break()

     if calendarios:
          for calendario in calendarios or []:
               img_base64 = calendario.get('img')
               if img_base64:
                    try:
                         doc.add_heading('Gráfico de distribución por fecha:', level= 2)
                         doc.add_paragraph("")
                         insertar_imagen(img_base64, 3, doc)
                    except:
                         doc.add_paragraph('Error al cargar calendario')

     if graphic:
        for g in graphic or []:
            img = g.get('img')
            if img:
                try:
                    titulo = 'Gráfico de distribución horaria:'
                    doc.add_heading(titulo, level=2)
                    doc.add_paragraph("")
                    insertar_imagen(img, 3, doc)
                except:
                    doc.add_paragraph('Error al agregar gráfico')

     if tabla_img:
        try:
            if isinstance(tabla_img, bytes):
                tabla_img = tabla_img.decode('utf-8')

            tabla_img = tabla_img.strip().replace('\n', '')
            titulo = 'Tabla de datos:'

            doc.add_heading(titulo, level=2)
            doc.add_paragraph("")
            insertar_imagen(tabla_img, 3, doc)

        except:
            doc.add_paragraph('Error en la tabla de datos')
     if horas:
          try:
               doc.add_heading('Rango horario crítico:', level= 2)
               doc.add_paragraph( horas)
          except:
               doc.add_paragraph("No hay rango horario crítico")
     

     response = HttpResponse(
          content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
     )

     fechaHora = datetime.datetime.now(timezone.utc)
     if task_id:
          db.collection('Reportes').document(task_id).delete()

     response['Content-Disposition'] = f'attachment; filename=Analisis_de_eventos_{now_str}.docx'
     doc.save(response)
     request.session.pop("task_id", None)

     return response