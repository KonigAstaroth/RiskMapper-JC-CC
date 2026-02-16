import re
from docx import Document



def markdown_to_docx(texto_markdown, doc: Document):
    lineas = texto_markdown.split("\n")
    i = 0

    while i < len(lineas):
        linea = lineas[i].rstrip()

        #Encabezados
        match = re.match(r'^(#{1,4})\s+(.*)', linea)
        if match:
            nivel = len(match.group(1))
            texto = match.group(2).strip()
            doc.add_heading(texto, level=nivel)
            i += 1
            continue

        # Tablas - markdown
        if linea.strip().startswith("|"):
            filas = []

            while i < len(lineas) and lineas[i].strip().startswith("|"):
                filas.append(lineas[i].strip())
                i += 1

            # eliminar linea separadora tipo |---|
            if len(filas) > 1 and re.match(r'^\|\s*[-:]+\s*(\|\s*[-:]+\s*)+\|?$', filas[1]):
                filas.pop(1)

            data = []
            for fila in filas:
                columnas = [col.strip() for col in fila.strip("|").split("|")]
                data.append(columnas)

            if data:
                tabla = doc.add_table(rows=len(data), cols=len(data[0]))
                tabla.style = "Table Grid"

                for row_idx, fila in enumerate(data):
                    for col_idx, celda in enumerate(fila):
                        tabla.rows[row_idx].cells[col_idx].text = celda

            continue

        # Listas
        if re.match(r'^[-*]\s+', linea):
            texto = re.sub(r'^[-*]\s+', '', linea)
            add_bold_paragraph(doc, texto, "List Bullet" )
            i += 1
            continue

        if re.match(r'^\d+\.\s+', linea):
            texto = re.sub(r'^\d+\.\s+', '', linea)
            doc.add_paragraph(texto, style="List Number")
            i += 1
            continue

        # Parrafo normal
        if linea.strip():
            doc.add_paragraph(linea.strip())

        i += 1


def add_bold_paragraph(doc, texto, estilo=None):
    parrafo = doc.add_paragraph(style=estilo)

    partes = re.split(r'(\*\*.*?\*\*)', texto)

    for parte in partes:
        if parte.startswith("**") and parte.endswith("**"):
            run = parrafo.add_run(parte[2:-2])
            run.bold = True
        else:
            parrafo.add_run(parte)

    return parrafo