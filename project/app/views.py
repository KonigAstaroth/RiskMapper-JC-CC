
import requests
from django.shortcuts import render,redirect
from django.conf import settings
from django.contrib import messages
from firebase_admin import firestore, auth
from datetime import timedelta
import urllib.parse
import datetime
from datetime import timezone
import pandas as pd
import json
from django.utils.safestring import mark_safe
from django.utils import timezone as dj_timezone
from google.cloud.firestore import FieldFilter
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib
import calendar
matplotlib.use('agg')
import numpy as np
import io
import base64
from time import time
from geopy.geocoders import GoogleV3
from collections import defaultdict
from django.core.cache import cache
from .utils import sendEmail
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO
from django.http import HttpResponse
from itsdangerous import URLSafeTimedSerializer, SignatureExpired, BadSignature
from collections import Counter
from openai import OpenAI
import os,psutil
import re
from bs4 import BeautifulSoup
from docx.enum.text import WD_ALIGN_PARAGRAPH
import gc
import tracemalloc
import stripe






db = firestore.client()
CACHE_KEY_MARKERS = 'firebase_markers'
CACHE_KEY_LAST_UPDATE = 'firebase_markers_last_update'

FIREBASE_AUTH_URL = f"https://identitytoolkit.googleapis.com/v1/accounts:signInWithPassword?key={settings.FIREBASE_API_KEY}"
stripe.api_key = settings.STRIPE_SECRET_KEY

def signup(request):
     if request.method == 'POST':
          name = request.POST.get('name')
          lastName = request.POST.get('lastName')
          email = request.POST.get('correo_personal')
          password = request.POST.get('clave_segura')
          repassword = request.POST.get('clave_repetida')
          client = request.POST.get('client')
          try:
               if name and lastName and email and password == repassword:
                    user = auth.create_user(email=email, password=password)
                    db.collection("Usuarios").document(user.uid).set({
                    "email": email,
                    "name": name,
                    "lastname": lastName,
                    "privileges":False,
                    "lastAccess": None,
                    "client_type": client
                    })
                    request.session['email_usr'] = email
                    return redirect('subs')
               elif repassword != password:
                    error_message = "La contraseña no coicide"
                    return redirect(f"/signup?error={urllib.parse.quote(error_message)}") 
               else:
                    error_message = "Faltan campos por ser llenados"
                    return redirect(f"/signup?error={urllib.parse.quote(error_message)}") 
          except Exception as e:
               return redirect(f"/signup?error={urllib.parse.quote(e)}")
                  
     error = request.GET.get("error")
     return render(request, 'signup.html', {"error": error})

def subscriptions(request):
     # if  not all( k in request.session for k in ['name_usr', 'lastname_usr', 'email_usr', 'password_usr', 'client_usr' ]):
     #      return redirect('signup')
     if request.method == "POST":
          plan = request.POST.get('plan')
          request.session['plan'] = plan
          return create_stripe_user(request)

     return render(request, 'selectSub.html')

def create_stripe_user(request):
     email = request.session.get('email_usr')
     plan = request.session.get('plan')
     customer= None

     if not email or not plan:
          return HttpResponse("Faltan datos en la sesión", status=400)
     
     try:
          
          user = auth.get_user_by_email(email)
     except:
          return HttpResponse('Usuario no autenticado', status=401)


     user_doc = db.collection('Usuarios').document(user.uid)
     user_data = user_doc.get().to_dict()

     customer_id = user_data.get('stripe_customer_id')

     if not 'stripe_customer_id' in user_data:
          customer = stripe.Customer.create(
               email= user_data.get('email'),
               name= user_data.get('name') + " " + user_data.get('lastname')
          )
          customer_id = customer.id
          user_doc.update({'stripe_customer_id': customer_id})


     
     prices = {
          'esencial': 'price_1RnkWOPF5qcM1JsRRk5zLJlj',
          'premium': 'price_1RnkZ2PF5qcM1JsRiwtjRrw4',
          'profesional': 'price_1RnkbEPF5qcM1JsRSjBsHD9l'
     }

     if plan not in prices:
        return HttpResponse("Plan inválido", status=400)

     checkout_session = stripe.checkout.Session.create(
          customer= customer_id,
          payment_method_types=['card'],
          line_items=[{
               'price': prices[plan],
               'quantity': 1
          }],
          mode= 'subscription',
          success_url='http://127.0.0.1:8000/success?session_id={CHECKOUT_SESSION_ID}',
          cancel_url='http://127.0.0.1:8000/cancel',
          metadata={
                'firebase_uid': user.uid,
                'plan': plan
          }

     )
     return redirect(checkout_session.url)

def success(request):
     # plan = request.session.get('plan')
     # email = request.session.get('email_usr')
     # try:
          
     #      user = auth.get_user_by_email(email)
     # except:
     #      return HttpResponse('Usuario no autenticado', status=401)
     
     # if plan not in ['esencial', 'premium', 'profesional']:
     #      return HttpResponse("Plan inválido", status=400)
     
     # PLAN_CONFIG ={
     #      'esencial': {'requests': 10, 'events': 100},
     #      'premium': {'requests': 50, 'events': 200},
     #      'profesional': {'requests': 80, 'events': 2000}
     # }

     # config=PLAN_CONFIG[plan]
     # sub_type = plan
     # requests = config['requests']
     # events = config['events']

     # user_doc = db.collection('Usuarios').document(user.uid)

     # now = datetime.datetime.now(timezone.utc)
     # endSub = now + datetime.timedelta(days=30)
     # end_sub = datetime.datetime.combine(endSub.date(), time(23,59), tzinfo=timezone.utc)
     # try:
     #      user_doc.update({
     #           'sub_type': sub_type,
     #           'requests': requests,
     #           'events': events,
     #           'start_sub': now,
     #           'last_sub': now,
     #           'end_sub': end_sub,
     #           'status': "active"

     #      })
     # except Exception as e:
     #      return HttpResponse('Usuario no encontrado', status=401)

     return render(request, 'success.html')

def login(request):
    sessionCookie = request.COOKIES.get('session')
    if sessionCookie:
         decoded_claims = auth.verify_session_cookie(sessionCookie, check_revoked=True)
         uid = decoded_claims["uid"]
         db.collection('Usuarios').document(uid).update({'lastAccess': datetime.datetime.now(timezone.utc)})
         return redirect ("main")

    if request.method == 'POST':
        recaptcha_response = request.POST.get('g-recaptcha-response')

        data = {
            'secret': settings.RECAPTCHA_SECRET_KEY,
            'response': recaptcha_response
        }
        verify_url = 'https://www.google.com/recaptcha/api/siteverify'
        result = requests.post(verify_url, data=data).json()

        if not result.get('success'):
            messages.error(request, "Se debe de validar el CAPTCHA")
            return redirect('/')
        
        email = request.POST["email"]
        password = request.POST["password"]
        remember = request.POST.get("remember")
             
        login = {"email": email, "password": password, "returnSecureToken": True}
        auth_response = requests.post(FIREBASE_AUTH_URL, json=login)
        
        if auth_response.status_code == 200:
                user_data = auth_response.json()
                id_token = user_data["idToken"]
                if remember == "checked":
                     expires_in = timedelta(days=14)
                else: 
                    expires_in = timedelta(days=1)

                try:
                    session_cookie = auth.create_session_cookie(id_token, expires_in=expires_in)
                    response_redirect = redirect("main")
                    response_redirect.set_cookie(
                        key='session',
                        value=session_cookie,
                        max_age=expires_in.total_seconds(),
                        httponly=True,
                        secure=True  # Cambiar a True en producción con HTTPS
                    )
                    decoded_claims = auth.verify_session_cookie(session_cookie)
                    uid = decoded_claims["uid"]
                    db.collection('Usuarios').document(uid).update({'lastAccess': datetime.datetime.now(timezone.utc)})
                    
                    return response_redirect
                except Exception as e:
                    messages.error(request, "Error al crear la cookie de sesión:", str(e))
                    return redirect('/')
        else:
             messages.error(request, "Correo o contraseña incorrectos") 
             return redirect('/')
    return render(request, 'login.html')

def policy (request):
    return render(request, 'policy.html')

def generate_token(email, uid):
     s = URLSafeTimedSerializer(settings.SECRET_KEY)
     return s.dumps(email, salt="recover-pass")

def forgotpass (request):
    if request.method == "POST":
          try:
               
               email = request.POST.get('email')
               user = auth.get_user_by_email(email)
               if user:
                    token = generate_token(email, user.uid)
                    link = request.build_absolute_uri(f"https://riskmapper-jc-cc.onrender.com/recoverPass/{token}/")
                    sendEmail(email, link)
                    success_message = "Correo para restablecer contraseña ha sido enviado"
                    return redirect(f"/forgotPassword?success={urllib.parse.quote(success_message)}")
                    
          except Exception as e:
               error_message = f"{e}"
               return redirect(f"/forgotPassword?error={urllib.parse.quote(error_message)}")
                  
    error = request.GET.get("error")
    success = request.GET.get("success")
    return render(request, 'forgotPass.html', {'error':error, 'success': success})

def recoverPass (request, token):
     s = URLSafeTimedSerializer(settings.SECRET_KEY)
     
     try:
          email = s.loads(token, salt="recover-pass", max_age = 900)
     except SignatureExpired:
          error_message = "El enlace ha expirado. Solicita uno nuevo."
          return redirect(f"/recoverPass/{token}?error={urllib.parse.quote(error_message)}") 
     except BadSignature:
          error_message = "El enlace es inválido."
          return redirect(f"/recoverPass/{token}?error={urllib.parse.quote(error_message)}") 
     
     if not email:
          return redirect('login')
     contra = request.POST.get('password')
     repassword = request.POST.get('repassword')
     
     if request.method == 'POST':
          
          if not contra or not repassword:
               error_message = "Faltan campos por llenar"
               return redirect(f"/recoverPass/{token}?error={urllib.parse.quote(error_message)}") 
          if contra != repassword:
               error_message = "Las contraseñas no coinciden"
               return redirect(f"/recoverPass/{token}?error={urllib.parse.quote(error_message)}")  
          
          try:
               
               user = auth.get_user_by_email(email)
               auth.update_user(user.uid, password=contra)
               success_message = "Contraseña restablecida"
               return redirect(f"/recoverPass/{token}?success={urllib.parse.quote(success_message)}")
          except Exception as e:
               error_message = str(e)
               return redirect(f"/recoverPass/{token}?error={urllib.parse.quote(error_message)}") 
          

     error = request.GET.get("error")
     success = request.GET.get('success')
     return render(request, 'recoverPass.html', {'error': error, 'success':success})

     

def time_to_num(time_str):
     try:
          hh, mm, ss = map(int, time_str.strip().split(':'))
          return ss +60*(mm + 60*hh)
     except Exception as e:
          print(f"Error al convertir '{time_str}': {e}")
          return None
     
def getRange(eventos):
     horas = []

     for evento in eventos:
        fecha = evento.get('FechaHoraHecho')
        if not fecha:
            print("Evento sin fecha:", evento)
            continue

        
        fecha_dt = None
        if isinstance(fecha, str):
            try:
                fecha_dt = datetime.fromisoformat(fecha)
            except Exception as e:
                print("Error al convertir string a datetime:", e)
                continue
        elif hasattr(fecha, 'to_datetime'):
            try:
                fecha_dt = fecha.to_datetime()
            except Exception as e:
                print("Error al usar to_datetime():", e)
                continue
        elif isinstance(fecha, datetime.datetime):
            fecha_dt = fecha
        else:
            continue

        if fecha_dt.tzinfo is not None:
            fecha_local = dj_timezone.localtime(fecha_dt)
        else:
            fecha_local = fecha_dt  

        horas.append(fecha_local.hour)
     

     if not horas:
        return None, None

     ctr = Counter(horas)
     hora_critica, cantidad = ctr.most_common(1)[0]
     return hora_critica, cantidad

def main (request):
     sessionCookie = request.COOKIES.get('session')
     priv = getPrivileges(request)

     if not sessionCookie:
          return redirect ("login")
     try:
          decoded_claims = auth.verify_session_cookie(sessionCookie, check_revoked=True)
          uid = decoded_claims["uid"]
     except:
          return redirect("login")
     
     doc_ref = db.collection("Usuarios").document(uid)
     doc = doc_ref.get()

     if doc.exists:
          name = doc.to_dict().get("name")


     usuarios = getUsers()


     # datos necesarios para google maps
     last_update = cache.get(CACHE_KEY_LAST_UPDATE)
     markers_cached = cache.get(CACHE_KEY_MARKERS)

          
     ref = db.collection('Eventos')
     
     query_has_update = False
     
     
     if last_update:
          changes_query = ref.where('updatedAt', '>', last_update)
          changes = changes_query.limit(1).get()
          query_has_update = len(changes) > 0
     else:
          query_has_update =True
     
     map_config = {
        'center': {'lat': 19.42847, 'lng': -99.12766},
        'zoom': 6
     }

     if query_has_update:
          data = ref.limit(500).get() or []
          list_markers = []
          max_update = last_update

          for doc in data:
               valor = doc.to_dict()
               update_doc = valor.get('updatedAt')
               if update_doc:
                    if not max_update or update_doc > max_update:
                         max_update = update_doc

               
               fecha_obj = valor.get('FechaHoraHecho')
               fecha_str = ''
               if fecha_obj:
                    fecha_local = dj_timezone.localtime(fecha_obj)
                    fecha_str = fecha_local.strftime('%d/%m/%Y, %H:%M:%S')
               marker = {
                    'lat': valor.get('latitud'),
                    'lng': valor.get('longitud'),
                    'Categoria': valor.get('Categoria'),
                    'icono': valor.get('icono'),
                    'delito': valor.get('Delito'),
                    'fecha': fecha_str,
                    'calle': valor.get('Calle_hechos'),
                    'colonia': valor.get('ColoniaHechos'),
                    'estado': valor.get('Estado_hechos'),
                    'municipio': valor.get('Municipio_hechos'),
                    'descripcion': valor.get('Descripcion')
               }
               list_markers.append(marker)
          cache.set(CACHE_KEY_MARKERS, list_markers, None)
          cache.set(CACHE_KEY_LAST_UPDATE, max_update, None)
     else:
          list_markers = markers_cached or []

     markers_json = mark_safe(json.dumps(list_markers))

     #Filtrado de datos
     graphic = request.session.get('graphic')
     calendarios = request.session.get('calendarios', [])
     
     hour_txt = request.session.get('hour_txt', None)
     AiText = request.session.get('AiText', None)
     lugar = request.session.get('lugar')
     tabla = request.session.get('tabla_base64', None)
     

     if request.method != 'POST':
          if not request.session.get('desde_busqueda'):
               request.session.pop('desde_busqueda', None)
               
          else:
               if not request.session.get('ready_to_export'):
                    request.session.pop('graphic', None)
                    request.session.pop('calendarios', None)
                    request.session.pop('lugar', None)
                    request.session.pop('hour_txt', None)
                    request.session.pop('AiText', None)
                    request.session.pop('tabla_base64', None)
                    map_config = request.session.pop('map_config', {
                         'center': {'lat': 19.42847, 'lng': -99.12766},
                         'zoom': 6
                         })
          if 'map_config' in request.session:
               map_config = request.session['map_config']
               
     
     
     

     lista_delitos= [
          {'valor': 'AMENAZAS', 'nombre': 'Amenazas'},
          {'valor': 'ROBO A NEGOCIO', 'nombre': 'Robo a negocio'},
          {'valor': 'HOMICIDIO DOLOSO', 'nombre': 'Homicidio doloso'},
          {'valor': 'FEMINICIDIO', 'nombre': 'Feminicidio'},
          {'valor': 'SECUUESTRO', 'nombre': 'Secuestro'},
          {'valor': 'TRATA DE PERSONAS', 'nombre': 'Trata de personas'},
          {'valor': 'ROBO A TRANSEÚNTE', 'nombre': 'Robo a transeúnte'},
          {'valor': 'EXTORSIÓN', 'nombre': 'Extorsión'},
          {'valor': 'ROBO A CASA HABITACIÓN', 'nombre': 'Robo a casa habitación'},
          {'valor': 'VIOLACIÓN', 'nombre': 'Violación'},
          {'valor': 'NARCOMENUDEO', 'nombre': 'Narcomenudeo'},
          {'valor': 'CATEGORIA DE BAJO IMPACTO', 'nombre': 'Delito de bajo impacto'},
          {'valor': 'ARMA DE FUEGO', 'nombre': 'Lesión con arma de fuego'},
          {'valor': 'ROBO DE ACCESORIOS DE AUTO', 'nombre': 'Robo de accesorios de auto'},
          {'valor': 'ROBO A CUENTAHABIENTE SALIENDO DEL CAJERO CON VIOLENCIA', 'nombre': 'Robo a cuentahabiente'},
          {'valor': 'ROBO DE VEHÍCULO', 'nombre': 'Robo de vehículo'},
          {'valor': 'ROBO A PASAJERO A BORDO DE MICROBUS', 'nombre': 'Robo en microbús'},
          {'valor': 'ROBO A REPARTIDOR', 'nombre': 'Robo a repartidor'},
          {'valor': 'ROBO A PASAJERO A BORDO DEL METRO', 'nombre': 'Robo en metro'},
          {'valor': 'LESIONES DOLOSAS POR DISPARO DE ARMA DE FUEGO', 'nombre': 'Lesiones por arma de fuego'},
          {'valor': 'HECHO NO DELICTIVO', 'nombre': 'Hecho no delictivo'},
          {'valor': 'ROBO A PASAJERO A BORDO DE TAXI CON VIOLENCIA', 'nombre': 'Robo en taxi'},
          {'valor': 'ROBO A TRANSPORTISTA', 'nombre': 'Robo a transportista'},
     ]

     color_delitos = [
          {'valor': 'AMENAZAS', 'nombre': 'Amenazas', 'color': '#8A2BE2'},
          {'valor': 'ROBO A NEGOCIO', 'nombre': 'Robo a negocio', 'color': '#FF69B4'},
          {'valor': 'HOMICIDIO DOLOSO', 'nombre': 'Homicidio doloso', 'color': '#4B0082'},
          {'valor': 'FEMINICIDIO', 'nombre': 'Feminicidio', 'color': '#DC143C'},
          {'valor': 'SECUUESTRO', 'nombre': 'Secuestro', 'color': '#FF1493'},
          {'valor': 'TRATA DE PERSONAS', 'nombre': 'Trata de personas', 'color': '#B22222'},
          {'valor': 'ROBO A TRANSEÚNTE', 'nombre': 'Robo a transeúnte', 'color': '#00FF00'},
          {'valor': 'EXTORSIÓN', 'nombre': 'Extorsión', 'color': '#DAA520'},
          {'valor': 'ROBO A CASA HABITACIÓN', 'nombre': 'Robo a casa habitación', 'color': '#800080'},
          {'valor': 'VIOLACIÓN', 'nombre': 'Violación', 'color': '#8B0000'},
          {'valor': 'NARCOMENUDEO', 'nombre': 'Narcomenudeo', 'color': '#006400'},
          {'valor': 'CATEGORIA DE BAJO IMPACTO', 'nombre': 'Delito de bajo impacto', 'color': '#A0A0A0'},
          {'valor': 'ARMA DE FUEGO', 'nombre': 'Lesión con arma de fuego', 'color': '#00008B'},
          {'valor': 'ROBO DE ACCESORIOS DE AUTO', 'nombre': 'Robo de accesorios de auto', 'color': '#A0522D'},
          {'valor': 'ROBO A CUENTAHABIENTE SALIENDO DEL CAJERO CON VIOLENCIA', 'nombre': 'Robo a cuentahabiente', 'color': '#FF0000'},
          {'valor': 'ROBO DE VEHÍCULO', 'nombre': 'Robo de vehículo', 'color': '#FFA500'},
          {'valor': 'ROBO A PASAJERO A BORDO DE MICROBUS', 'nombre': 'Robo en microbús', 'color': '#FFD700'},
          {'valor': 'ROBO A REPARTIDOR', 'nombre': 'Robo a repartidor', 'color': '#ADFF2F'},
          {'valor': 'ROBO A PASAJERO A BORDO DEL METRO', 'nombre': 'Robo en metro', 'color': '#00FFFF'},
          {'valor': 'LESIONES DOLOSAS POR DISPARO DE ARMA DE FUEGO', 'nombre': 'Lesiones por arma de fuego', 'color': '#00008B'},
          {'valor': 'HECHO NO DELICTIVO', 'nombre': 'Hecho no delictivo', 'color': '#D3D3D3'},
          {'valor': 'ROBO A PASAJERO A BORDO DE TAXI CON VIOLENCIA', 'nombre': 'Robo en taxi', 'color': '#40E0D0'},
          {'valor': 'ROBO A TRANSPORTISTA', 'nombre': 'Robo a transportista', 'color': '#0000FF'},
     ]


     
     if request.method == 'POST' and 'buscar' in request.POST :
          # process = psutil.Process(os.getpid())
          # mem_inicio = process.memory_info().rss / 1024**2
          # cpu_inicio = psutil.cpu_percent(interval=0.1)

          # print(f"[ANTES] Memoria: {mem_inicio:.2f} MB, CPU: {cpu_inicio}%")

          # tracemalloc.start() 
          filters = {}
          filtersAi = {}
          direccion = ""
          graphic = []
          calendarios = []
          banner = []
          
          str_startDate = None
          str_endDate = None
          descripcion_cliente = ""
          
          
          calle = request.POST.get('calle')
          colonia = request.POST.get('colonia')
          municipio = request.POST.get('municipio')
          estado = request.POST.get('estado')
          startDate_str = request.POST.get('startDate')
          endDate_str = request.POST.get('endDate')
          delitos_select = request.POST.getlist('delitos')
          descripcion_cliente = request.POST.get('descripcion_cliente')


          if calle:
               filters['Calle_hechos'] = calle
               filtersAi['Calle'] = calle
               direccion += calle + ', '
               banner.append(calle)
          if colonia:
               filters['ColoniaHechos'] = colonia
               filtersAi['Colonia']= colonia
               direccion += colonia + ', '
               banner.append(colonia)
          if municipio:
               filters['Municipio_hechos'] = municipio
               filtersAi['Municipio'] = municipio
               direccion += municipio + ', '
               banner.append(municipio)
          if estado:
               filters['Estado_hechos'] = estado
               filtersAi['Estado'] = estado
               direccion += estado + ', '
               banner.append(estado)

          

          
          map_config = getLatLng(direccion)
          lugar = ', '.join(f"{k}" for k in banner)
          
          
 
          if startDate_str and endDate_str:
               startDate = datetime.datetime.strptime(startDate_str, "%Y-%m-%d").replace(tzinfo=timezone.utc)
               endDate = datetime.datetime.strptime(endDate_str, "%Y-%m-%d").replace(tzinfo=timezone.utc) + timedelta(days=1)
               str_startDate = startDate.strftime("%Y-%m-%d")
               str_endDate = endDate.strftime("%Y-%m-%d")
               filters['startDate']=startDate
               filters['endDate']=endDate
          else:
               pass

          now = datetime.datetime.now(timezone.utc)
          AiText = genAI(filtersAi, str_startDate,str_endDate, descripcion_cliente, now)
          

          if not (('startDate' in filters and 'endDate' in filters) or any(k in filters for k in ['Municipio_hechos', 'Estado_hechos', 'Calle_hechos', 'ColoniaHechos'])):
               request.session['graphic'] = []
               request.session['calendarios'] = []
               request.session['tabla_base64'] = None
               request.session['ready_to_export'] = True
               request.session['hour_txt'] = "No se encontraron eventos para graficar."
               request.session['AiText'] = mark_safe(AiText)
               request.session['lugar'] = lugar
               request.session['map_config'] = map_config
               return redirect("main")
               
          

          query_ref = ref

          if filters:
               for campo, valor in filters.items():
                    if campo == "startDate":
                         query_ref = query_ref.where(filter=FieldFilter("FechaHoraHecho", '>=', valor))
                    elif campo == "endDate":
                         query_ref = query_ref.where(filter=FieldFilter("FechaHoraHecho", '<=', valor))
                    else:
                         if isinstance(valor, str):
                              valor = valor.strip() 
                         query_ref = query_ref.where(filter=FieldFilter(campo, '==', valor))
          if delitos_select:
               if len(delitos_select)<= 10:
                    query_ref = query_ref.where(filter=FieldFilter("Categoria", "in", delitos_select))
               else:
                    error_message = "Demasiados delitos seleccionados para aplicar filtro 'in'"
                    return redirect(f"/main?error={urllib.parse.quote(error_message)}")
               
          preview = list(query_ref.limit(1).stream())

          if not preview:
               request.session['graphic'] = []
               request.session['calendarios'] = []
               request.session['tabla_base64'] = None
               request.session['ready_to_export'] = True
               request.session['hour_txt'] = "No se encontraron eventos para graficar."
               request.session['AiText'] = mark_safe(AiText)
               request.session['lugar'] = lugar
               request.session['map_config'] = map_config
               return redirect("main")

          # mem_despues = process.memory_info().rss / 1024**2
          # cpu_despues = psutil.cpu_percent(interval=0.1)

          # print(f"[DESPUÉS] Memoria: {mem_despues:.2f} MB, CPU: {cpu_despues}%")
          # print(f"[DIFERENCIA] Memoria usada durante proceso: {mem_despues - mem_inicio:.2f} MB")

          

          eventos_por_mes = defaultdict(list)
          graficos_por_mes = defaultdict(list)
          cat_color = []

          

          resultados = list(query_ref.stream())

          eventos_lista=[doc.to_dict() for doc in resultados]
          conteo_delitos = defaultdict(int)
          

          hora_critica, cantidad = getRange(eventos_lista)
          
          if hora_critica is not None:
               if cantidad > 1:
                    hour_txt = f"Entre las {hora_critica}:00 y las {hora_critica+1}:00 horas se registró {cantidad} eventos, lo que destaca este intervalo como un posible punto de riesgo."
               elif cantidad == 1:
                     hour_txt = f"Entre las {hora_critica}:00 y las {hora_critica+1}:00 horas se registró {cantidad} evento, lo que destaca este intervalo como un posible punto de riesgo."
          else:
               hour_txt = "No hay eventos para calcular rango horario crítico."

          
          for  eventos in eventos_lista:
               date_obj = eventos.get('FechaHoraHecho')
               if not date_obj:
                    continue
               fecha = date_obj
               categorie = eventos.get('Categoria')
               

               match = next((item for item in color_delitos if item['valor'] == categorie), None)
               if match:
                    nombre = match['nombre']
                    color = match['color']
               else:
                    nombre = 'Otro'
                    color = 'gray'

               conteo_delitos[nombre] += 1

               if not any(c['nombre'] == nombre for c in cat_color):
                    cat_color.append({'nombre': nombre, 'color': color})

               if fecha:
                    if isinstance(fecha,str):
                         fecha = datetime.fromisoformat(fecha)
                    elif hasattr(fecha, 'to_datetime'):
                         fecha = fecha.to_datetime()

                    year = fecha.year
                    month = fecha.month
                    day = fecha.day
               

               eventos_por_mes[(year, month)].append((day, categorie))
               

               if date_obj:
                    hora_local = dj_timezone.localtime(date_obj)
                    hora_str = hora_local.strftime('%H:%M:%S').strip()
                    hora_str = ''.join(c for c in hora_str if c.isdigit() or c == ':')
                    segundos = time_to_num(hora_str)
                    if segundos is not None:
                         angulo = (segundos/86400) * 2 * np.pi
                         dia = hora_local.day
                         graficos_por_mes[(year, month)].append((angulo, dia, categorie))
                    
          for (year, month), day in eventos_por_mes.items():
                    if not day:
                         continue
                    else:
                         imagen_base64  = generateCalendar(year, month, day)
                         if imagen_base64:
                              calendarios.append({
                                   'img': imagen_base64 
                              })
                         else:
                              error_message = "No se pudo el calendario"
                              return redirect(f"/main?error={urllib.parse.quote(error_message)}")
                    puntos_mes = graficos_por_mes.get((year,month), [])
                    if not puntos_mes:
                         continue
                    else:
                         if puntos_mes:
                              graphic_img = genGraph(puntos_mes)
                              graphic.append({
                                   'img': graphic_img
                              })
                         else:
                              error_message = "No se pudo la grafica"
                              return redirect(f"/main?error={urllib.parse.quote(error_message)}")  
                         
          cat_color_cuenta = []
          for item in cat_color:
               nombre = item['nombre']
               color = item['color']
               cuenta = conteo_delitos.get(nombre, 0)
               cat_color_cuenta.append({'nombre': nombre, 'color': color, 'cuenta': cuenta})
          
          tabla = genDataImg(cat_color_cuenta)
          
     

          # snapshot = tracemalloc.take_snapshot()
          # top_stats = snapshot.statistics('lineno')

          # print("[TRACEMALLOC] Top 10 líneas con mayor consumo de memoria:")
          # for stat in top_stats[:10]:
          #      print(stat)
          request.session['graphic'] = graphic
          request.session['calendarios'] = calendarios
          request.session['lugar'] = lugar
          request.session['hour_txt'] = hour_txt
          request.session['desde_busqueda'] = True
          request.session['AiText'] = mark_safe(AiText)
          request.session['map_config'] = map_config
          request.session['tabla_base64'] = tabla
          request.session['ready_to_export'] = True
          return redirect('main') 
          
     error = request.GET.get("error")

     
               

     context = {
          'name': name,
          'priv': priv,
          'usuarios': usuarios,
          'google_maps_api_key': settings.GOOGLE_MAPS_KEY,
          'markers': markers_json,
          'graphic': graphic,
          'calendarios': calendarios,
          'timestamp': int(time()),
          'lugar': lugar,
          'hour_txt': hour_txt,
          'AiText': AiText,
          'map_config_json': json.dumps(map_config),
          'error': error,
          'lista_delitos': lista_delitos,
          'tabla_base64': tabla
          
     }

    
     return render (request, 'main.html', context)

def genDataImg(cat_color_cuenta):
     n = len(cat_color_cuenta)
     radio = 0.4                  
     espacio = radio * 5          

     alto_total = n * espacio  

     alto_fig = max(alto_total * 0.4, 1)
     fig, ax = plt.subplots(figsize=(6, alto_fig * 0.4))  
     ax.set_aspect('equal')
     ax.axis('off')

     for i, item in enumerate(cat_color_cuenta):
        nombre = item['nombre']
        color = item['color']
        y = (n - i - 1) * espacio

        
        icon = mpatches.Circle((0.5, y), radio, color=color)
        ax.add_patch(icon)

        
        cuenta = item.get('cuenta', 0)
        texto = f"{nombre.upper()}: {cuenta}"
        ax.text(1.2, y, texto, va='center', fontsize=12)

    
     ax.set_xlim(-0.2, 5)
     ax.set_ylim(-radio - 0.5, alto_total - espacio + radio + 0.2)

     plt.tight_layout(pad=0.5)

     buffer = io.BytesIO()
     plt.savefig(buffer, format='png', bbox_inches = 'tight')
     plt.close(fig)
     ax.clear()
     buffer.seek(0)
     img_png = buffer.getvalue()
     tabla = base64.b64encode(img_png).decode('utf-8')
     buffer.close()
     del fig, ax, buffer, img_png
     gc.collect()

     return tabla

def genGraph(puntos):
          if not puntos:
               return None
          
          colores_cat ={
          'DELITO DE BAJO IMPACTO': '#A0A0A0',  # gris
          'ROBO A CUENTAHABIENTE SALIENDO DEL CAJERO CON VIOLENCIA': '#FF0000',  # rojo
          'ROBO DE VEHÍCULO CON Y SIN VIOLENCIA': '#FFA500',  # naranja
          'VIOLACIÓN': '#8B0000',  # rojo oscuro
          'ROBO A PASAJERO A BORDO DE MICROBUS CON Y SIN VIOLENCIA': '#FFD700',  # dorado
          'ROBO A REPARTIDOR CON Y SIN VIOLENCIA': '#ADFF2F',  # verde lima
          'ROBO A PASAJERO A BORDO DEL METRO CON Y SIN VIOLENCIA': '#00FFFF',  # cian
          'LESIONES DOLOSAS POR DISPARO DE ARMA DE FUEGO': '#00008B',  # azul oscuro
          'ROBO A NEGOCIO CON VIOLENCIA': '#FF69B4',  # rosa fuerte
          'HECHO NO DELICTIVO': '#D3D3D3',  # gris claro
          'ROBO A TRANSEUNTE EN VÍA PÚBLICA CON Y SIN VIOLENCIA': '#00FF00',  # verde
          'ROBO A PASAJERO A BORDO DE TAXI CON VIOLENCIA': '#40E0D0',  # turquesa
          'HOMICIDIO DOLOSO': '#4B0082',  # índigo
          'ROBO A CASA HABITACIÓN CON VIOLENCIA': '#800080',  # morado
          'SECUESTRO': '#FF1493',  # rosa mexicano
          'ROBO A TRANSPORTISTA CON Y SIN VIOLENCIA': '#0000FF',
          'AMENAZAS': '#8A2BE2',  # azul violeta
          'ROBO A NEGOCIO': '#FF69B4',  # misma que "ROBO A NEGOCIO CON VIOLENCIA"
          'FEMINICIDIO': '#DC143C',  # carmesí
          'SECUUESTRO': '#FF1493',  # igual que "SECUESTRO"
          'TRATA DE PERSONAS': '#B22222',  # rojo fuego
          'ROBO A TRANSEÚNTE': '#00FF00',  # igual que transeúnte con violencia
          'EXTORSIÓN': '#DAA520',  # dorado oscuro
          'ROBO A CASA HABITACIÓN': '#800080',  # igual que "CON VIOLENCIA"
          'NARCOMENUDEO': '#006400',  # verde oscuro
          'ARMA DE FUEGO': '#00008B',  # igual que lesiones por arma
          'ROBO DE ACCESORIOS DE AUTO': '#A0522D',  # marrón
          'ROBO A PASAJERO A BORDO DE MICROBUS': '#FFD700',  # igual que con violencia
          'ROBO A REPARTIDOR': '#ADFF2F',  # igual que con violencia
          'ROBO A PASAJERO A BORDO DEL METRO': '#00FFFF',  # igual que con violencia
          'ROBO A TRANSPORTISTA': '#0000FF',
          }

          angulos =[]
          radios =[]
          colores = []

          for angulo, radio, categoria in puntos:
               angulos.append(angulo)
               radios.append(radio)
               colores.append(colores_cat.get(categoria, 'gray'))

          fig =  plt.figure(figsize=(6,6))
          ax = plt.subplot(111, polar=True)
          sc = ax.scatter(angulos, radios, color=colores, s=80, alpha=0.75)
          ax.set_theta_direction(-1)
          ax.set_theta_offset(np.pi/2)

          horas =[f"{h:02d}:00" for h in range(24)]
          ticks = [(h/24) * 2 * np.pi for h in range(24)]
          ax.set_xticks(ticks)
          ax.set_xticklabels(horas, fontsize=8)
          ax.set_rlabel_position(135) 
          ax.set_ylim(0,31)
          
          
          buffer = io.BytesIO()
          plt.savefig(buffer, format='png')
          buffer.seek(0)
          img_png = buffer.getvalue()
          graphic = base64.b64encode(img_png).decode('utf-8')
          plt.close(fig)
          ax.clear()
          buffer.close()
          del fig, ax, buffer, img_png
          gc.collect()
          return graphic

def getLatLng(direccion):
     try:
          geolocator = GoogleV3(api_key=settings.GOOGLE_MAPS_KEY)
          location = geolocator.geocode(direccion)
          if location:
               lat = location.latitude
               lng = location.longitude
          else:
               lat = None
               lng = None
          map_config = {
               'center': {'lat': float(lat) if lat is not None else 19.42847, 
                         'lng': float(lng) if lng is not None else -99.12766
               },
               'zoom': 14 if lat and lng else 6
          }
          return map_config
     except Exception as e:
          print(str(e))
          return {
            'center': {'lat': 19.42847, 'lng': -99.12766},
            'zoom': 6
        }
     
def loadOsintDate(name = "osintDate.txt" ):
     ruta = os.path.join(settings.BASE_DIR, 'app','prompts', name)
     with open (ruta, 'r', encoding='utf-8')as f:
          return f.read()
     
 

def genAI(filters,start,end, descripcion_cliente,now):
     client = OpenAI(api_key=settings.OPENAI_API_KEY)
     
     lugar = ', '.join(f"{k}:{v}" for k,v in filters.items()) if filters else "No especificado"
     template = loadOsintDate()
     content= template.format(start=start, end = end, lugar = lugar, descripcion_cliente = descripcion_cliente, now = now)
     completion =client.chat.completions.create(
          model='gpt-4.1-mini',
          store = True,
          messages=[{'role': 'user', 'content': content}]
     )

     
     text = completion.choices[0].message.content.strip()
     
     return cleanAnswer(text)


def cleanAnswer(texto):
    texto = texto.strip()
    texto = re.sub(r'\n?[^.\n]*\?$', '', texto).strip()

    patrones_finales = [
        r'Si desea.*?$',
        r'¿Desea.*?$',
        r'¿Le gustaría.*?$',
        r'¿Quiere que.*?$',
    ]
    for patron in patrones_finales:
        texto = re.sub(patron, '', texto, flags=re.IGNORECASE).strip()

    
    texto = re.sub(r'^#{1,6}\s*(.*)', r'<h3>\1</h3>', texto, flags=re.MULTILINE)

    
    texto = re.sub(r'^[\-\*]\s+(.*)', r'<li>\1</li>', texto, flags=re.MULTILINE)
    if '<li>' in texto:
        texto = '<ul>' + texto + '</ul>'

    texto = re.sub(r'^\d+\.\s+(.*)', r'<li>\1</li>', texto, flags=re.MULTILINE)
    if re.search(r'<li>.*</li>', texto) and not texto.startswith('<ul>'):
        texto = '<ol>' + texto + '</ol>'

    # Negritas y cursivas
    texto = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', texto)
    texto = re.sub(r'\*(.*?)\*', r'<em>\1</em>', texto)

    
    def procesar_tabla(match):
        tabla = match.group(0)
        filas = tabla.strip().split('\n')
        if len(filas) < 2:
            return tabla

        cabecera = filas[0]
        cuerpo = filas[2:] if re.match(r'^\|[-\s|]+\|$', filas[1]) else filas[1:]

        def fila_a_html(fila, tag='td'):
            celdas = [f'<{tag} style="border: 1px solid #000; padding: 5px;">{c.strip()}</{tag}>' for c in fila.strip('|').split('|')]
            return '<tr>' + ''.join(celdas) + '</tr>'

        tabla_html = '<table style="border-collapse: collapse; width: 100%;">'
        tabla_html += fila_a_html(cabecera, tag='th')
        for fila in cuerpo:
            if fila.strip():
                tabla_html += fila_a_html(fila)
        tabla_html += '</table>'

        return tabla_html

    
    texto = re.sub(
        r'((?:\|.*\|\n)+\|[-\s|]+\|\n(?:\|.*\|\n?)*)',
        procesar_tabla,
        texto,
        flags=re.MULTILINE
    )

    
    bloques = texto.split('\n\n')
    html_parts = []
    for b in bloques:
        b = b.strip()
        if not b:
            continue
        if '<table' in b:
            html_parts.append(b)
        else:
            html_parts.append(f'<p>{b}</p>')

    return ''.join(html_parts)




def exportarDocx(request):
     graphic = request.session.get('graphic')
     calendarios = request.session.get('calendarios', [])
     horas = request.session.get('hour_txt')
     AiText = request.session.get('AiText')
     text_html = AiText
     tabla = request.session.get('tabla_base64')

     doc = Document()
     
     doc.add_heading('Análisis de eventos', 0)
     if text_html:
        soup = BeautifulSoup(text_html, 'html.parser')
        for elemento in soup.contents:
            if elemento.name == 'p':
                doc.add_paragraph(elemento.get_text())
            elif elemento.name == 'h3':
                doc.add_heading(elemento.get_text(), level=2)
            elif elemento.name == 'ul':
                for li in elemento.find_all('li'):
                    doc.add_paragraph('• ' + li.get_text(), style='List Bullet')
            elif elemento.name == 'ol':
                for li in elemento.find_all('li'):
                    doc.add_paragraph(li.get_text(), style='List Number')
            elif elemento.name == 'table':
                filas = elemento.find_all('tr')
                if filas:
                    num_cols = len(filas[0].find_all(['td', 'th']))
                    tabla = doc.add_table(rows=1, cols=num_cols)
                    tabla.style = 'Table Grid'

                    # Encabezado
                    hdr_cells = tabla.rows[0].cells
                    for idx, th in enumerate(filas[0].find_all(['td', 'th'])):
                        hdr_cells[idx].text = th.get_text().strip()
                        for paragraph in hdr_cells[idx].paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(11)

                    # Cuerpo de la tabla
                    for fila in filas[1:]:
                        celdas = fila.find_all(['td', 'th'])
                        row_cells = tabla.add_row().cells
                        for idx, celda in enumerate(celdas):
                            row_cells[idx].text = celda.get_text().strip()
                            for paragraph in row_cells[idx].paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                    # Ajustar ancho de columnas (opcional)
                    ancho_columna = Inches(1.5)
                    for col_idx in range(num_cols):
                        for row in tabla.rows:
                            row.cells[col_idx].width = ancho_columna

     for calendario in calendarios:
          img_base64 = calendario.get('img')
          if img_base64:
               if img_base64.startswith('data:image'):
                    img_base64 = img_base64.split(',')[1]
               try:
                    calendario_data = base64.b64decode(img_base64)
                    calendario_stream = BytesIO(calendario_data)
                    doc.add_heading('Gráfico de distribución por fecha:', level= 2)
                    doc.add_picture(calendario_stream, width=Inches(3))  
               except:
                    doc.add_paragraph('Error al cargar calendario')

     if graphic :
          for g in graphic:
               img = g.get('img')
               if img.startswith('data:image'):
                    img = img.split(',')[1]
               try:
                    graphic_data = base64.b64decode(img)
                    graphic_stream = BytesIO(graphic_data)
                    doc.add_heading('Gráfico de distribución horaria:', level= 2)
                    doc.add_picture(graphic_stream, width=Inches(4))  
               except:
                    doc.add_paragraph('Error al agregar calendario')

     if tabla:
          
          try:
               if isinstance(tabla, str) and tabla.startswith('data:image'):
                    tabla = tabla.split(',')[1]
               tabla_data = base64.b64decode(tabla)
               tabla_stream = BytesIO(tabla_data)
               doc.add_heading('Tabla de datos:', level= 2)
               doc.add_picture(tabla_stream, width=Inches(3))
          except:
               doc.add_paragraph('Error en la tabla de datos')
     if horas:
          try:
               doc.add_heading('Rango horario crítico:', level= 2)
               doc.add_paragraph( horas)
          except:
               doc.add_paragraph("No hay rango horario crítico")
     

     response = HttpResponse(
          content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
     )

     fechaHora = datetime.datetime.now(timezone.utc)

     response['Content-Disposition'] = f'attachment; filename =Analisis_de_eventos{fechaHora}.docx'
     doc.save(response)
     request.session.pop('graphic', None)
     request.session.pop('calendarios', None)
     request.session.pop('hour_txt', None)
     request.session.pop('AiText', None)
     request.session.pop('ready_to_export', None)
     request.session.pop('tabla_base64', None)

     return response

def generateCalendar( year: int, month: int, eventos: list[tuple]) -> str:
     cal = calendar.monthcalendar(year, month)
     fig, ax = plt.subplots(figsize=(5,5))
     ax.set_axis_off()
     ax.set_title(calendar.month_name[month] + f" {year}", fontsize=20)

     dias_semana = ["Lun", "Mar", "Mié", "Jue", "Vie", "Sáb", "Dom"]

     colores_cat ={
          'DELITO DE BAJO IMPACTO': '#A0A0A0',  # gris
          'ROBO A CUENTAHABIENTE SALIENDO DEL CAJERO CON VIOLENCIA': '#FF0000',  # rojo
          'ROBO DE VEHÍCULO CON Y SIN VIOLENCIA': '#FFA500',  # naranja
          'VIOLACIÓN': '#8B0000',  # rojo oscuro
          'ROBO A PASAJERO A BORDO DE MICROBUS CON Y SIN VIOLENCIA': '#FFD700',  # dorado
          'ROBO A REPARTIDOR CON Y SIN VIOLENCIA': '#ADFF2F',  # verde lima
          'ROBO A PASAJERO A BORDO DEL METRO CON Y SIN VIOLENCIA': '#00FFFF',  # cian
          'LESIONES DOLOSAS POR DISPARO DE ARMA DE FUEGO': '#00008B',  # azul oscuro
          'ROBO A NEGOCIO CON VIOLENCIA': '#FF69B4',  # rosa fuerte
          'HECHO NO DELICTIVO': '#D3D3D3',  # gris claro
          'ROBO A TRANSEUNTE EN VÍA PÚBLICA CON Y SIN VIOLENCIA': '#00FF00',  # verde
          'ROBO A PASAJERO A BORDO DE TAXI CON VIOLENCIA': '#40E0D0',  # turquesa
          'HOMICIDIO DOLOSO': '#4B0082',  # índigo
          'ROBO A CASA HABITACIÓN CON VIOLENCIA': '#800080',  # morado
          'SECUESTRO': '#FF1493',  # rosa mexicano
          'ROBO A TRANSPORTISTA CON Y SIN VIOLENCIA': '#0000FF',
          'AMENAZAS': '#8A2BE2',  # azul violeta
          'ROBO A NEGOCIO': '#FF69B4',  # misma que "ROBO A NEGOCIO CON VIOLENCIA"
          'FEMINICIDIO': '#DC143C',  # carmesí
          'SECUUESTRO': '#FF1493',  # igual que "SECUESTRO"
          'TRATA DE PERSONAS': '#B22222',  # rojo fuego
          'ROBO A TRANSEÚNTE': '#00FF00',  # igual que transeúnte con violencia
          'EXTORSIÓN': '#DAA520',  # dorado oscuro
          'ROBO A CASA HABITACIÓN': '#800080',  # igual que "CON VIOLENCIA"
          'NARCOMENUDEO': '#006400',  # verde oscuro
          'ARMA DE FUEGO': '#00008B',  # igual que lesiones por arma
          'ROBO DE ACCESORIOS DE AUTO': '#A0522D',  # marrón
          'ROBO A PASAJERO A BORDO DE MICROBUS': '#FFD700',  # igual que con violencia
          'ROBO A REPARTIDOR': '#ADFF2F',  # igual que con violencia
          'ROBO A PASAJERO A BORDO DEL METRO': '#00FFFF',  # igual que con violencia
          'ROBO A TRANSPORTISTA': '#0000FF',
    }

     for i, dias in enumerate(dias_semana):
          ax.text(i+0.5, len(cal), dias, ha='center', va= 'center', fontsize=12, weight='bold')
     

     for fila, semana in enumerate(cal):
          for columna, dia in enumerate(semana):
               if dia != 0:
                    y = len(cal) - fila - 0.5
                    x = columna + 0.5
                    ax.text(x, y + 0.3, str(dia), ha='center', va='top', fontsize=10)
                    for (dia_evento, categoria) in eventos:
                         if dia_evento == dia:
                              color = colores_cat.get(categoria, 'gray')
                              circle = plt.Circle((x, y -0.2), 0.15, color= color)
                              ax.add_patch(circle) 

     plt.xlim(0, 7)
     plt.ylim(0, len(cal) + 1)
     plt.tight_layout()

     buffer = io.BytesIO()
     plt.savefig(buffer, format='png')
     buffer.seek(0)
     imagen_png = buffer.getvalue()
     imagen_base64 = base64.b64encode(imagen_png).decode('utf-8')
     buffer.close()
     
     plt.close()
     return imagen_base64 

def getPrivileges(request):
      sessionCookie = request.COOKIES.get('session')
      decoded_claims = auth.verify_session_cookie(sessionCookie, check_revoked=True)
      uid = decoded_claims["uid"]
      doc_ref = db.collection("Usuarios").document(uid)
      doc = doc_ref.get()
      if doc.exists:
           return doc.to_dict().get("privileges",False)
      else:
           return False


def logout (request):
     response = redirect("login")
     response.delete_cookie('session')
     request.session.flush()
     
     return response

def add (request):
     priv = getPrivileges(request)
     sessionCookie = request.COOKIES.get('session')
     
     try:
          decoded_claims = auth.verify_session_cookie(sessionCookie, check_revoked=True)
          uid = decoded_claims["uid"]
     except:
          return redirect("login")
     
     priv = getPrivileges(request)
     if not priv:
               return redirect("main")
     
     if not sessionCookie:
          return redirect ("login")
     
     if request.method == "POST":
          
        email = request.POST.get("email")
        password = request.POST.get("password")
        name = request.POST.get("name")
        lastname = request.POST.get("lastname")
        privileges = request.POST.get('privileges')

        if name and lastname and password and email and privileges:
             
             try:
                  user = auth.create_user(email=email, password=password)
                  db.collection("Usuarios").document(user.uid).set({
                    "email": email,
                    "name": name,
                    "lastname": lastname,
                    "privileges":privileges,
                    "lastAccess": None
                    
                    })
                  success_message = "Usuario agregado correctamente"
                  return redirect(f"/add?success={urllib.parse.quote(success_message)}")
             except Exception as e:
                  error_message = str(e)
                  return redirect(f"/add?error={urllib.parse.quote(error_message)}")
                  
        else:
             error_message = "Faltan campos por ser llenados"
             return redirect(f"/add?error={urllib.parse.quote(error_message)}")
                  
     success = request.GET.get("success")
     error = request.GET.get("error")
     return render (request, "addUser.html", {"success": success, "error": error,'priv': priv})

def manage_user(request):
     usuarios = getUsers()
     priv = getPrivileges(request)
     sessionCookie = request.COOKIES.get('session')
     
     try:
          decoded_claims = auth.verify_session_cookie(sessionCookie, check_revoked=True)
          uid = decoded_claims["uid"]
     except:
          return redirect("login")
     
     priv = getPrivileges(request)
     if not priv:
               return redirect("main")
     
     if not sessionCookie:
          return redirect ("login")
     return render(request, 'manageUser.html', {"usuarios": usuarios, "priv":priv})

def getUsers():
     docs = db.collection("Usuarios").stream()
     listUsers=[]

     for doc in docs:
          datos = doc.to_dict()
          datos['id'] = doc.id
          listUsers.append(datos)

     return listUsers

def editUser(request, id):
     
     if request.method == 'POST':
        uid = id
        updates = {}

        if name := request.POST.get('name'):
            updates['name'] = name

        if lastname := request.POST.get('lastname'):
            updates['lastname'] = lastname

        if email := request.POST.get('email'):
            updates['email'] = email
            if uid:
                try:
                    auth.update_user(uid, email=email)
                except Exception as e:
                    print("Error actualizando email:", e)

        if password := request.POST.get('password'):
            if uid:
                try:
                    auth.update_user(uid, password=password)
                except Exception as e:
                    print("Error actualizando contraseña:", e)

        if privileges := request.POST.get('privileges'):
            updates['privileges'] = privileges == 'true'

        if updates:
            db.collection('Usuarios').document(id).update(updates)
     return redirect('manageUser')
               

def deleteUser(request, id):
     if request.method == 'POST':
          doc_ref = db.collection('Usuarios').document(id)
          doc = doc_ref.get()
          if doc.exists:
               uid = id
               doc_ref.delete()
               if uid:
                    try:
                         auth.delete_user(uid)
                    except auth.UserNotFoundError:
                         pass
          
     return redirect('manageUser')

def display_users(request):
     usuarios = getUsers()
     priv = getPrivileges(request)
     sessionCookie = request.COOKIES.get('session')
     
     try:
          decoded_claims = auth.verify_session_cookie(sessionCookie, check_revoked=True)
          uid = decoded_claims["uid"]
     except:
          return redirect("login")
     
     priv = getPrivileges(request)
     if not priv:
               return redirect("main")
     
     if not sessionCookie:
          return redirect ("login")
     return (request,{"usuarios": usuarios, "priv":priv})

def loadFiles(request):
    sessionCookie = request.COOKIES.get('session')
    priv = getPrivileges(request)

    try:
        decoded_claims = auth.verify_session_cookie(sessionCookie, check_revoked=True)
        uid = decoded_claims["uid"]
    except Exception as e:
        print("Error autenticando:", e)
        return redirect("login")

    if not sessionCookie:
        return redirect("login")
    usuarios = getUsers()
    geolocator = GoogleV3(api_key=settings.GOOGLE_MAPS_KEY)

    if request.method == "POST":
        if 'archivo' in request.FILES:
            try:
                excel_file = request.FILES['archivo']
                df = pd.read_excel(excel_file, engine='openpyxl')

                def convertir_fecha(fecha):
                    if isinstance(fecha, pd.Timestamp):
                         return fecha.to_pydatetime()
                    if isinstance(fecha, datetime.datetime):
                         return fecha
                    if isinstance(fecha, str):
                         fecha = fecha.strip()
                         formatos_validos = [
                              "%d/%m/%Y",    
                              "%Y-%m-%d",      
                              "%-d/%-m/%Y",    
                              "%#d/%#m/%Y"     
                         ]
                         for fmt in formatos_validos:
                              try:
                                   return datetime.datetime.strptime(fecha, fmt)
                              except:
                                   continue
                    return None

                def convertir_hora(h):
                    if isinstance(h, str):
                         for fmt in ("%H:%M:%S", "%H:%M"):
                              try:
                                   return datetime.datetime.strptime(h, fmt).time()
                              except:
                                   continue
                    if isinstance(h, datetime.time):
                         return h
                    return None

                if "FechaHecho" in df.columns:
                    df['FechaHecho'] = pd.to_datetime(df['FechaHecho'], dayfirst=True, errors='coerce')
                if "HoraHecho" in df.columns:
                    df["HoraHecho"] = df["HoraHecho"].apply(convertir_hora)

                # Combinar fecha + hora en un solo datetime (Muy util para consultas mas delante)
                def combinar_fecha_hora(fecha, hora):
                    if isinstance(fecha, datetime.datetime) and isinstance(hora, datetime.time):
                        return datetime.datetime.combine(fecha.date(), hora)
                    return None

                df["FechaHoraHecho"] = df.apply(
                    lambda row: combinar_fecha_hora(row.get("FechaHecho"), row.get("HoraHecho")),
                    axis=1
                )

                df.drop(columns=["FechaHecho", "HoraHecho"], inplace=True, errors='ignore')

               

                datos = df.where(pd.notnull(df), None).to_dict(orient='records')
                campos_ubicacion=['latitud', 'longitud', 'Estado_hechos', 'Municipio_hechos', 'ColoniaHechos', 'Calle_hechos' ]

                for evento in datos:
                    try:
                        if location_check(evento, campos_ubicacion):
                              pass
                        else:
                         if check_valid_value(evento.get('latitud')) and check_valid_value(evento.get('longitud')):
                                   calle_dir = evento.get('Calle_hechos')
                                   col_dir = evento.get('ColoniaHechos')
                                   estado_dir = evento.get('Estado_hechos')
                                   municipio_dir = evento.get('Municipio_hechos')
                                   if check_valid_value(estado_dir) and check_valid_value(municipio_dir):
                                        try:
                                             if 'Calle_hechos2' in df.columns and check_valid_value(evento.get('Calle_hechos2')):
                                                  calle_dir_dos = evento.get('Calle_hechos2')
                                                  direccion = f"{calle_dir}, {calle_dir_dos}, {col_dir}, {municipio_dir}, {estado_dir}"
                                             else:
                                                  direccion = f"{calle_dir}, {col_dir}, {municipio_dir}, {estado_dir}"
                                             location = geolocator.geocode(direccion)
                                             if location:
                                                  evento['latitud'] = location.latitude
                                                  evento['longitud'] = location.longitude
                                        except:
                                             error_message = "No se encontraron las coordenadas: Revise la dirección"
                                             return redirect(f"/loadFiles?error={urllib.parse.quote(error_message)}")
                                   else:
                                        
                                        municipio_geo, estado_geo = getEstadoMunicipio(location)

                                        if municipio_geo:
                                             evento['Municipio_hechos'] = municipio_geo
                                        if estado_geo:
                                             evento['Estado_hechos'] = estado_geo
                         if check_valid_value(evento.get('Estado_hechos')) and check_valid_value(evento.get('Municipio_hechos')):
                                   
                                   try:
                                        latitud = evento.get('latitud', '')
                                        longitud = evento.get('longitud', '')
                                        if latitud and longitud:
                                             location = geolocator.reverse((latitud,longitud))
                                             if location:
                                                  municipio_geo, estado_geo = getEstadoMunicipio(location)
                                                  if municipio_geo:
                                                       evento['Municipio_hechos'] = municipio_geo
                                                  if estado_geo:
                                                       evento['Estado_hechos'] = estado_geo
                                   except:
                                        error_message = "No se encontró el estado y/o municipio: Error en las coordenadas"
                                        return redirect(f"/loadFiles?error={urllib.parse.quote(error_message)}")
                        
                        
                        
                             

                        categoria = evento.get('Categoria', '')
                        icono = None
                        if isinstance(categoria, str):
                            categoria_lower = categoria.lower()
                            if 'amenazas' in categoria_lower:
                                icono = 'amenazas'
                            elif 'robo a negocio' in categoria_lower:
                                icono = 'robonegocio'
                            elif 'homicidio doloso' in categoria_lower:
                                 icono = 'homicidiodoloso'
                            elif 'feminicidio' in categoria_lower:
                                 icono = 'feminicidio'
                            elif 'secuestro' in categoria_lower:
                                 icono = 'secuestro'
                            elif 'trata de personas' in categoria_lower:
                                 icono = 'tratapersonas'
                            elif 'robo a transeúnte' in categoria_lower:
                                 icono = 'robotranseunte'
                            elif 'extorsión' in categoria_lower:
                                 icono = 'extorsion'
                            elif 'robo a casa habitación' in categoria_lower:
                                 icono = 'robocasa'
                            elif 'violación' in categoria_lower:
                                 icono = 'violacion'
                            elif 'narcomenudeo' in categoria_lower:
                                 icono = 'narcomenudeo'
                            elif 'categoria de bajo impacto' in categoria_lower or 'delito de bajo impacto' in categoria_lower:
                                 icono = "bajoimpacto"
                            elif 'arma de fuego' in categoria_lower:
                                 icono = 'armafuego'
                            elif 'robo de accesorios de auto' in categoria_lower:
                                 icono= 'robovehiculo'
                            elif 'robo a cuentahabiente saliendo del cajero con violencia' in categoria_lower:
                                   icono = 'robocuentahabiente'
                            elif 'robo de vehículo' in categoria_lower:
                                   icono = 'robovehiculo'
                            elif 'robo a pasajero a bordo de microbus' in categoria_lower:
                                   icono = 'robomicrobus'
                            elif 'robo a repartidor' in categoria_lower:
                                   icono = 'roborepartidor'
                            elif 'robo a pasajero a bordo del metro' in categoria_lower:
                                   icono = 'robometro'
                            elif 'lesiones dolosas por disparo de arma de fuego' in categoria_lower:
                                   icono = 'armafuego'
                            elif 'hecho no delictivo' in categoria_lower:
                                   icono = 'nodelito'
                            elif 'robo a pasajero a bordo de taxi con violencia' in categoria_lower:
                                   icono = 'robotaxi'
                            elif 'robo a transportista' in categoria_lower:
                                   icono = 'robotransportista'
                            else:
                                 icono = 'default'

                        
                        evento['icono'] = icono
                        evento['updatedAt'] = datetime.datetime.now(datetime.timezone.utc)
                        categoria = evento.get('Categoria', '').upper()
                        evento['Categoria'] = categoria
                        
                        db.collection('Eventos').add(evento)

                    except Exception as e:
                        error_message = "Error subiendo evento a Firestore"
                        return redirect(f"/loadFiles?error={urllib.parse.quote(error_message)}")
                        
                        

                success_message = "La carga ha sido exitosa"
                return redirect(f"/loadFiles?success={urllib.parse.quote(success_message)}")

            except Exception as e:
                error_message = "Error general"
                return redirect(f"/loadFiles?error={urllib.parse.quote(error_message)}")

        # Carga manual
        else:
            calle = request.POST.get("calle")
            colonia = request.POST.get("colonia")
            estado = request.POST.get("estado")
            municipio = request.POST.get("municipio")
            delito = request.POST.get("crime")
            fechaValue = request.POST.get("FechaHoraHecho")
            icon = request.POST.get("icons")
            categoria = request.POST.get("categ")
            lat = request.POST.get('lat')
            lng = request.POST.get('lng')
            descripcion = request.POST.get('descripcion')

            if not descripcion:
                 descripcion = ""
                 
            
            crime_str = str(categoria).upper()
          # Este segmento es sobre los iconos
            try:
                icono = None
                if 'amenazas' in icon:
                         icono = 'amenazas'
                elif 'robo a negocio' in icon:
                         icono = 'robonegocio'
                elif 'homicidio doloso' in icon:
                         icono = 'homicidiodoloso'
                elif 'feminicidio' in icon:
                         icono = 'feminicidio'
                elif 'secuestro' in icon:
                         icono = 'secuestro'
                elif 'trata de personas' in icon:
                         icono = 'tratapersonas'
                elif 'robo a transeúnte' in icon:
                         icono = 'robotranseunte'
                elif 'extorsión' in icon:
                         icono = 'extorsion'
                elif 'robo a casa habitación' in icon:
                         icono = 'robocasa'
                elif 'violación' in icon:
                         icono = 'violacion'
                elif 'narcomenudeo' in icon:
                         icono = 'narcomenudeo'
                elif 'categoria de bajo impacto' in icon or 'delito de bajo impacto' in icon:
                         icono = "bajoimpacto"
                elif 'arma de fuego' in icon:
                         icono = 'armafuego'
                elif 'robo de accesorios de auto' in icon:
                         icono= 'robovehiculo'
                elif 'robo a cuentahabiente saliendo del cajero con violencia' in icon:
                         icono = 'robocuentahabiente'
                elif 'robo de vehículo' in icon:
                         icono = 'robovehiculo'
                elif 'robo a pasajero a bordo de microbus' in icon:
                         icono = 'robomicrobus'
                elif 'robo a repartidor' in icon:
                         icono = 'roborepartidor'
                elif 'robo a pasajero a bordo del metro' in icon:
                         icono = 'robometro'
                elif 'lesiones dolosas por disparo de arma de fuego' in icon:
                         icono = 'armafuego'
                elif 'hecho no delictivo' in icon:
                         icono = 'nodelito'
                elif 'robo a pasajero a bordo de taxi con violencia' in icon:
                         icono = 'robotaxi'
                elif 'robo a transportista' in icon:
                         icono = 'robotransportista'
                else:
                     icono = 'default'
                
                      
            except Exception as e:
                 print(e)

            timestamp = None

            if fechaValue:
                 try:
                    if 'T' in fechaValue:
                        dt= datetime.datetime.fromisoformat(fechaValue) 
                    else:
                        dt=datetime.datetime.strptime(fechaValue, "%Y-%m-%d %H:%M:%S")

                    timestamp = dj_timezone.make_aware(dt)
                        
                 except Exception as e:
                      return render(request, 'loadFiles.HTML', {
                           'error': f'Error al convertir la fecha: {e}'
                      })
                 
                      
                           
          
            if ((calle and colonia and estado and municipio) or (lat and lng)) and fechaValue and delito and icono:
                now = datetime.datetime.now(datetime.timezone.utc)
                try:
                    if ((not calle and not colonia and not estado and not municipio) and lat and lng):
                         
                         location = geolocator.reverse((lat,lng))
                         if location and location.raw:
                              componentes = location.raw.get('address_components', [])
                              calle_geo = colonia_geo = municipio_geo = estado_geo = None
                              for comp in componentes:
                                   tipos = comp['types']
                                   if 'route' in tipos :
                                        calle_geo = comp['long_name']
                                   elif 'street_number' in tipos:
                                        numero_geo = comp['long_name']
                                   elif 'sublocality' in tipos or 'sublocality_level_1' in tipos:
                                        colonia_geo = comp['long_name']
                                   elif 'locality' in tipos:
                                        municipio_geo = comp['long_name']
                                   elif 'administrative_area_level_1' in tipos:
                                        estado_geo = comp['long_name']

                                   if calle_geo and numero_geo:
                                        calleNumero = f"{calle_geo} {numero_geo}"
                                   else:
                                        calleNumero = calle_geo

                              db.collection('Eventos').add({
                              "Calle_hechos": calleNumero,
                              "ColoniaHechos": colonia_geo,
                              "Estado_hechos": estado_geo,
                              "Delito": delito,
                              "FechaHoraHecho": timestamp,
                              "icono": icono,
                              "Municipio_hechos": municipio_geo,
                              "Categoria": crime_str,
                              "latitud": lat,
                              "longitud": lng,
                              'updatedAt': now,
                              'Descripcion': descripcion
                              })
                    if (( calle and colonia and  estado and  municipio) and (not lat and not lng)):
                         direccion=f"{calle}, {colonia}, {estado}, {municipio}"
                         location = geolocator.geocode(direccion)
                         lat_geo = location.latitude
                         lng_geo = location.longitude

                         db.collection('Eventos').add({
                               "Calle_hechos": calle,
                               "ColoniaHechos": colonia,
                               "Estado_hechos": estado,
                               "Delito": delito,
                               "FechaHoraHecho": timestamp,
                               "icono": icono,
                               "Municipio_hechos": municipio,
                               "Categoria": crime_str,
                               "latitud": lat_geo,
                               "longitud": lng_geo,
                               'updatedAt': now,
                               'Descripcion': descripcion
                               })
                         

                    elif calle and colonia and estado and municipio and lat and lng:
                     
                         db.collection('Eventos').add({
                              "Calle_hechos": calle,
                              "ColoniaHechos": colonia,
                              "Estado_hechos": estado,
                              "Delito": delito,
                              "FechaHoraHecho": timestamp,
                              "icono": icono,
                              "Municipio_hechos": municipio,
                              "Categoria": crime_str,
                              "latitud": lat,
                              "longitud": lng,
                              'updatedAt': now,
                              'Descripcion': descripcion
                              })
                    success_message = "Datos agregados exitosamente"
                    return redirect(f"/loadFiles?success={urllib.parse.quote(success_message)}")
                except:
                     error_message = "Error al subir datos"
                     return redirect(f"/loadFiles?error={urllib.parse.quote(error_message)}")
            else:
                 error_message = "Faltan datos por agregar"
                 return redirect(f"/loadFiles?error={urllib.parse.quote(error_message)}")
     
    success = request.GET.get("success")
    error = request.GET.get("error")
    return render(request, "loadFiles.html", {"error": error, 'success': success, 'usuarios': usuarios, 'priv': priv,})

def check_valid_value(valor):
     return not(
          valor is None or str(valor).strip().upper() in ['NA', 'N/A', '', 'None']
     )

def location_check(evento, campos):
     return all(evento.get(campo) not in ['', None, 'NA', 'N/a', 'n/a', 'N/A', 'na'] for campo in campos)

def getEstadoMunicipio(location):
     municipio = estado = None
     if location and location.raw.get('address_components', []):
          for comp in location.raw['address_components']:
               tipos = comp['types']
               if 'locality' in tipos:
                    municipio = comp['long_name']
               elif 'administrative_area_level_1' in tipos:
                    estado = comp['long_name']
     return municipio, estado

def library(request):
    sessionCookie = request.COOKIES.get('session')
    priv = getPrivileges(request)

    if not sessionCookie:
        return redirect("login")
    try:
        decoded_claims = auth.verify_session_cookie(sessionCookie, check_revoked=True)
        uid = decoded_claims["uid"]
    except:
        return redirect("login")

    usuarios = getUsers()
    ref = db.collection('Eventos')
    eventos = []
    filters = {}
    
    if request.method == 'POST':
        

        startDate_str = request.POST.get('startDate')
        endDate_str = request.POST.get('endDate')
        direccion = request.POST.get('direccion', '')
        search = request.POST.get('searchBy')
        categoria = request.POST.get('categoria')

        partes_direccion = [parte.strip() for parte in direccion.split(',') if parte.strip()]

        if search == "full":
            calle = partes_direccion[0] if len(partes_direccion) > 0 else None
            colonia = partes_direccion[1] if len(partes_direccion) > 1 else None
            municipio = partes_direccion[2] if len(partes_direccion) > 2 else None
            estado = partes_direccion[3] if len(partes_direccion) > 3 else None

            if calle:
                filters['Calle_hechos'] = calle
            if colonia:
                filters['ColoniaHechos'] = colonia
            if municipio:
                filters['Municipio_hechos'] = municipio
            if estado:
                filters['Estado_hechos'] = estado
        elif search == "estado":
            estado = partes_direccion[0] if len(partes_direccion) > 0 else None
            if estado:
                filters['Estado_hechos'] = estado
        elif search == "municipio":
            municipio = partes_direccion[0] if len(partes_direccion) > 0 else None
            if municipio:
                filters['Municipio_hechos'] = municipio
        elif search == "estadoMunicipio":
            municipio = partes_direccion[0] if len(partes_direccion) > 0 else None
            estado = partes_direccion[1] if len(partes_direccion) > 1 else None
            if municipio:
                filters['Municipio_hechos'] = municipio.strip()
            if estado:
                filters['Estado_hechos'] = estado.strip()

        
        if startDate_str and endDate_str:
            filters['startDate'] = startDate_str
            filters['endDate'] = endDate_str

        if categoria:
            filters['Categoria'] = categoria

        
    if filters:
          query_ref = ref

          startDate_str = filters.get('startDate')
          endDate_str = filters.get('endDate')

          
          if startDate_str and endDate_str:
               startDate = datetime.datetime.strptime(startDate_str, "%Y-%m-%d").replace(tzinfo=timezone.utc)
               endDate = datetime.datetime.strptime(endDate_str, "%Y-%m-%d").replace(tzinfo=timezone.utc) + timedelta(days=1)
               query_ref = query_ref.where(filter=FieldFilter("FechaHoraHecho", '>=', startDate))
               query_ref = query_ref.where(filter=FieldFilter("FechaHoraHecho", '<=', endDate))

          
          filters_sin_fechas = {k: v for k, v in filters.items() if k not in ['startDate', 'endDate']}

          
          for campo, valor in filters_sin_fechas.items():
               print(f"Campo: {campo}, Valor: {valor}")
               if isinstance(valor, str):
                    valor = valor.strip()
               query_ref = query_ref.where(filter=FieldFilter(campo, '==', valor))

          resultados = query_ref.stream()
          for doc in resultados:
               data = doc.to_dict()
               data['id'] = doc.id
               eventos.append(data)

    

    return render(request, 'library.html', {
        'usuarios': usuarios,
        'eventos': eventos,
        'priv': priv
    })

def edit_event(request, id):

     updates = {}
     

     if request.method == 'POST':
          
          if calle := request.POST.get('calle'):
               updates['Calle_hechos'] = calle
          if colonia := request.POST.get('colonia'):
               updates['ColoniaHechos'] = colonia
          if municipio:= request.POST.get('municipio'):
               updates['Municipio_hechos'] = municipio
          if estado := request.POST.get('estado'):
               updates['Estado_hechos'] = estado
          if icono := request.POST.get('icons'):
               if 'amenazas' in icono:
                         icon = 'amenazas'
               elif 'robo a negocio' in icono:
                         icon = 'robonegocio'
               elif 'homicidio doloso' in icono:
                         icon = 'homicidiodoloso'
               elif 'feminicidio' in icono:
                         icon = 'feminicidio'
               elif 'secuestro' in icono:
                         icon = 'secuestro'
               elif 'trata de personas' in icono:
                         icon = 'tratapersonas'
               elif 'robo a transeúnte' in icono:
                         icon = 'robotranseunte'
               elif 'extorsión' in icono:
                         icon = 'extorsion'
               elif 'robo a casa habitación' in icono:
                         icon = 'robocasa'
               elif 'violación' in icono:
                         icon = 'violacion'
               elif 'narcomenudeo' in icono:
                         icon = 'narcomenudeo'
               elif 'categoria de bajo impacto' in icono or 'delito de bajo impacto' in icono:
                         icon = "bajoimpacto"
               elif 'arma de fuego' in icono:
                         icon = 'armafuego'
               elif 'robo de accesorios de auto' in icono:
                         icon= 'robovehiculo'
               elif 'robo a cuentahabiente saliendo del cajero con violencia' in icono:
                         icon = 'robocuentahabiente'
               elif 'robo de vehículo' in icono:
                         icon = 'robovehiculo'
               elif 'robo a pasajero a bordo de microbus' in icono:
                         icon = 'robomicrobus'
               elif 'robo a repartidor' in icono:
                         icon = 'roborepartidor'
               elif 'robo a pasajero a bordo del metro' in icono:
                         icon = 'robometro'
               elif 'lesiones dolosas por disparo de arma de fuego' in icono:
                         icon = 'armafuego'
               elif 'hecho no delictivo' in icono:
                         icon = 'nodelito'
               elif 'robo a pasajero a bordo de taxi con violencia' in icono:
                         icon = 'robotaxi'
               elif 'robo a transportista' in icono:
                         icon = 'robotransportista'
               elif 'default' in icono:
                         icon = 'default'
               updates['icono'] = icon
          if FechaHoraHecho := request.POST.get('FechaHoraHecho'):
               if 'T' in FechaHoraHecho:
                    FechaHora= datetime.datetime.fromisoformat(FechaHoraHecho) 
               else:
                    FechaHora = datetime.datetime.strptime(FechaHoraHecho, "%Y-%m-%d %H:%M:%S")

               if dj_timezone.is_naive(FechaHora):
                    timestamp = dj_timezone.make_aware(FechaHora)
               else:
                    timestamp = FechaHora
               updates['FechaHoraHecho'] = timestamp
          if categoria := request.POST.get('categoria'):
                updates['Categoria'] = categoria
          if delito := request.POST.get('delito'):
                updates['Delito'] = delito
          if descripcion := request.POST.get('descripcion'):
                updates['Descripcion'] = descripcion
          

     if updates:
          try:
               updateNow = datetime.datetime.now(timezone.utc)
               updates['updatedAt'] = updateNow
               db.collection('Eventos').document(id).update(updates)
          except Exception as e:
               print(e)


     return library(request)

def deleteEvent(request, id):
      
     if request.method == 'POST':
          doc_ref = db.collection('Eventos').document(id)
          doc = doc_ref.get()
          try:
               if doc.exists:
                    doc_ref.delete()
                    cache.delete(CACHE_KEY_LAST_UPDATE)
                    cache.delete(CACHE_KEY_MARKERS)
          except Exception as e:
               print(e)
     
     return library(request)




# Create your views here.
